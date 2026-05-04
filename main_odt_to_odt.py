import os
import subprocess
import csv
from docx import Document

class Candidato:
    def __init__(self, nome, cognome, sesso, giorno, mese, anno, comune_nascita, provincia_nascita, cf, ciclo):
        self.CF = cf
        self.nome = nome
        self.cognome = cognome
        self.sesso = sesso
        self.data_nascita = giorno + "/" + mese + "/" + anno
        self.comune_nascita = comune_nascita
        self.provincia_nascita = provincia_nascita
        self.ciclo = ciclo
        self.title = ""

    def assignTitle(self, T):
        self.title = T

def candidati_nomi(candidati):
    result = ""
    for c in candidati:
        result += c.nome + " " + c.cognome + "\n"
    return result

def candidati_documenti(candidati):
    result = ""
    for c in candidati:
        if c.sesso == "m":
            result += "\n\n\nDott. "
        else:
            result += "\n\n\nDott.ssa "
        result += c.nome + " " + c.cognome + " identificat" + ("o" if c.sesso == "m" else "a") + " con il seguente documento "
        result += ".........................................\n\nrilasciato da " + ".........................................\n\n"
        result += "Firma ........................................."
    return result

def effify(non_f_str: str):
    return eval(f'f"""{non_f_str}"""')



def replace_placeholder(doc, placeholder_name, new_text):
    """
    Sostituisce il placeholder in tutto il documento DOCX (paragrafi e tabelle),
    preservando la formattazione il più possibile.
    """

    # Definiamo una funzione interna per processare un singolo paragrafo.
    # Questo ci serve perché i paragrafi si trovano sia nel corpo del doc che nelle tabelle.
    def process_paragraph(p):
        if placeholder_name in p.text:
            # Word spesso spezza i placeholder in più "runs" (es. {{ in uno e nome }} nell'altro).
            # Per risolvere, uniamo logicamente i run, facciamo il replace e mettiamo tutto nel primo.
            runs = p.runs
            if not runs:
                return

            # Ricostruiamo il testo completo del paragrafo
            full_text = "".join(run.text for run in runs)

            if placeholder_name in full_text:
                # Eseguiamo la sostituzione
                new_full_text = full_text.replace(placeholder_name, str(new_text))

                # Svuotiamo tutti i run tranne il primo
                # e impostiamo il testo finale nel primo run per mantenere lo stile iniziale
                runs[0].text = new_full_text
                for i in range(1, len(runs)):
                    runs[i].text = ""

    # 1. Scansione di tutti i paragrafi (inclusi i Titoli/Headers)
    for p in doc.paragraphs:
        process_paragraph(p)

    # 2. Scansione di tutte le tabelle
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                # Ogni cella contiene uno o più paragrafi
                for p in cell.paragraphs:
                    process_paragraph(p)


from docx import Document
from docxcompose.composer import Composer


def merge_docx_files(file_master, file_da_aggiungere, file_output):
    """
    Unisce due file docx mantenendo stili e struttura.
    Il secondo file inizierà su una nuova pagina.
    """
    # 1. Carichiamo il primo documento (che farà da base)
    master = Document(file_master)
    composer = Composer(master)

    # 2. Carichiamo il secondo documento
    doc_temp = Document(file_da_aggiungere)

    # 3. Aggiungiamo un salto pagina alla fine del master
    # così il secondo documento inizia su una pagina pulita
    master.add_page_break()

    # 4. Uniamo i documenti
    # Composer gestisce internamente la duplicazione degli stili e delle tabelle
    composer.append(doc_temp)

    # 5. Salviamo il risultato finale
    composer.save(file_output)
    print(f"Documenti uniti con successo in: {file_output}")

# --- COME USARLO ---
# unisci_odt_perfetto("candidato_1.odt", "candidato_2.odt", "verbale_unito.odt")

# --- COME USARLO ---
# unisci_odt_perfetto("candidato_1.odt", "candidato_2.odt", "verbale_unito.odt")

# Esempio d'uso
# merge_two_files("verbale1.odt", "verbale2.odt", "documento_finale.odt")

os.makedirs("output", exist_ok=True)

time_start = "9:00"

time_end = "17:00"

day = "26/1/2026"

decreto = "3625/2025 del 16/12/2025"

presidente = "Prof.ssa Barbara Re"
segretario = "Prof. Han Van Der Aa"
componente = "Prof.ssa Xixi Lu"

candidati = []
cycles = set()

with open('candidates.csv', encoding="utf-8") as csvfile:
    reader = csv.DictReader(csvfile)
    for row in reader:
        C = Candidato(row["nome"], row["cognome"], row["sesso"].lower(), row["data_nascita"].split("/")[0], row["data_nascita"].split("/")[1], row["data_nascita"].split("/")[2], row["luogo_nascita"].strip(), row["prov_nascita"].strip(), row["codice_fiscale"].strip(), int(row["ciclo_numero"]))
        C.title = row["titolo"].strip()
        candidati.append(C)
        cycles.add(row["ciclo_numero"])

candidati_names = candidati_nomi(candidati)
candidati_ids = candidati_documenti(candidati)

cycles = list(cycles)
cycles.sort()
cycles = ", ".join(cycles)

document = Document('verbale_esame_finale_1_PNRR.docx')

replace_placeholder(document, '{{presidente}}', presidente)
replace_placeholder(document, '{{componente}}', componente)
replace_placeholder(document, '{{segretario}}', segretario)
replace_placeholder(document, '{{decreto}}', decreto)
replace_placeholder(document, '{{day}}', day)
replace_placeholder(document, '{{time_start}}', time_start)
replace_placeholder(document, '{{time_end}}', time_end)
replace_placeholder(document, '{{ids}}', candidati_ids)
replace_placeholder(document, '{{candidates_all}}', candidati_names)



document.save('output/output.docx')

for candidateN, c in enumerate(candidati):

    name = c.nome
    surname = c.cognome
    title = c.title
    gender = c.sesso
    cycle = c.ciclo
    birthdate = c.data_nascita
    birthplace = c.comune_nascita
    province = c.provincia_nascita

    candidate_document = Document('verbale_esame_finale_2_PNRR.docx')

    replace_placeholder(candidate_document, '{{number}}', str(candidateN + 1))
    replace_placeholder(candidate_document, '{{day}}', day)
    replace_placeholder(candidate_document, '{{time_start}}', time_start)
    replace_placeholder(candidate_document, '{{name}}', name)
    replace_placeholder(candidate_document, '{{surname}}', surname)
    replace_placeholder(candidate_document, '{{title}}', title)
    replace_placeholder(candidate_document, '{{gender}}', gender)
    replace_placeholder(candidate_document, '{{cycle}}', cycle)
    replace_placeholder(candidate_document, '{{date_of_birth}}', birthdate)
    replace_placeholder(candidate_document, '{{place_of_birth}}', birthplace)
    replace_placeholder(candidate_document, '{{province}}', province)
    replace_placeholder(candidate_document, '{{presidente}}', presidente)
    replace_placeholder(candidate_document, '{{componente}}', componente)
    replace_placeholder(candidate_document, '{{segretario}}', segretario)

    candidate_document.save(f'output/candidate_{candidateN}.docx')

    merge_docx_files('output/output.docx', f'output/candidate_{candidateN}.docx', 'output/output.docx')


#     candidate = effify(attachmentN)
#     result += candidate
#

#
#     candidate = effify(attachmentB)
#     result += candidate
#
#     candidate = effify(attestation)
#     result += candidate
#
#     candidate = effify(privacy)
#     result += candidate
#
# result += effify(attachmentA)
